# -*- coding: utf-8 -*-
"""
Created on Sun Jul 24 21:11:15 2022

@author: grant
"""

import os
from docx import Document
from WhiskyFunctions.parseError import parseError

def addInfoHistory(Data, Template):
    
    inserted = False
    
    temp = os.getcwd()
    
    temp = temp + "\\Templates\\Award History Template\\" + Template
    
    doc = Document(temp)
    
    paragraphs = doc.paragraphs
    
    table = doc.tables[0]
    
    FullName = Data[0][0]
    
    IDList = []
    IDCounter = 0
    
    DateList = []
    DateCounter = 0
    
    QualNameList = []
    QualNameCounter = 0
    
    GradeList = []
    GradeCounter = 0
    
    for entry in Data:
        
        IDList.append(entry[1])
        DateList.append(entry[2])
        QualNameList.append(entry[3])
        GradeList.append(entry[4])
        
    # Cycles through each paragraph in document to check for the "Insert Name" tag     
    
    for para in paragraphs:
        

        if "Insert Name" in para.text:
            items = para.runs
            for line in items:
                if "Insert Name" in line.text:
                    line.text = str(FullName)
                    inserted = True

    # Cycles through each cell in a table to check for each insert tag, then adds data and increases counter where applicable

    for row in table.rows:
        items = row.cells
        for cell in items:
            
            if "Insert Certificate Number " + str(IDCounter) in cell.text:
                
                if IDCounter >= len(IDList):
                    cell.text = ""
                    IDCounter += 1
                else:    
                    cell.text = "Cert No: " + str(IDList[IDCounter])
                    IDCounter += 1
                    inserted = True
            
            if "Insert Qualification Name " + str(QualNameCounter) in cell.text:
                if QualNameCounter >= len(QualNameList):
                    cell.text = ""
                    QualNameCounter += 1
                else:
                    cell.text = "Qualification-Name: " + str(QualNameList[QualNameCounter])
                    QualNameCounter += 1
                    inserted = True
            
            if "Insert Grade " + str(GradeCounter) in cell.text:
                if GradeCounter >= len(GradeList):
                    cell.text = ""
                    GradeCounter += 1
                else:
                    cell.text = "Grade Achieved: " + str(GradeList[GradeCounter])
                    GradeCounter += 1
                    inserted = True
    
            if "Insert Date " + str(DateCounter) in cell.text:
                if DateCounter >= len(DateList):
                    cell.text = ""
                    DateCounter += 1
                else:
                    cell.text = "Date: " + str(DateList[DateCounter])
                    DateCounter += 1
                    inserted = True
            

    if inserted == False:
        return(parseError(2))
    try:
        doc.save(os.getcwd() + "\\Output\\" + str(FullName) +  ".docx")
    except Exception as e:
        return("There was an error with Certificate generation, do you have the certificate of " + str(FullName) + " already open?")
    
    return("Certificate creation successful, please see " + str(FullName) + ".docx for completed certificate")